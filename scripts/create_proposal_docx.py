#!/usr/bin/env python3
"""
Generate a professionally designed Word document for Poker Dream Business Proposal
Single delivery version - no phases
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Brand Colors
PRIMARY_COLOR = RGBColor(0x2C, 0x3E, 0x50)      # Dark Blue
SECONDARY_COLOR = RGBColor(0x34, 0x49, 0x5E)   # Slate
ACCENT_COLOR = RGBColor(0xC0, 0xA0, 0x62)      # Gold
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)        # Dark Gray
SUCCESS_COLOR = RGBColor(0x27, 0xAE, 0x60)     # Green
DANGER_COLOR = RGBColor(0xC0, 0x39, 0x2B)      # Red

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C0A062')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_styled_table(doc, data, header=True, col_widths=None):
    """Create a professionally styled table"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    if i == 0 and header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            if i == 0 and header:
                set_cell_shading(cell, '2C3E50')
            elif i % 2 == 0:
                set_cell_shading(cell, 'F8F9FA')

    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    return table

def add_section_header(doc, title):
    """Add a section header"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    add_horizontal_line(doc)

def add_subsection_header(doc, title):
    """Add a subsection header"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR

def add_h3_header(doc, title):
    """Add a h3 header"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR

def main():
    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ==================== COVER PAGE ====================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('POKER DREAM')
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Business Proposal')
    run.font.size = Pt(24)
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_data = [
        ['Version:', '2.0'],
        ['Date:', '5 December 2025'],
        ['Validity:', '30 days (until 5 January 2026)'],
    ]

    for label, value in info_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label} ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.color.rgb = TEXT_COLOR

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_section_header(doc, 'Table of Contents')

    toc_items = [
        '1. Executive Summary',
        '2. Goals & Outcomes',
        '3. What You Get',
        '4. Project Timeline (12 Weeks)',
        '5. Capacity Fit',
        '6. Your Responsibilities',
        '7. Maintenance & Support',
        '8. Pricing & Payment',
        '9. Data & Privacy',
        '10. Acceptance & Next Steps',
    ]

    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_page_break()

    # ==================== 1. EXECUTIVE SUMMARY ====================
    add_section_header(doc, '1. Executive Summary')

    p = doc.add_paragraph()
    run = p.add_run('Poker Dream is an innovative mobile application designed specifically for poker lovers in Malaysia and around the world. The app offers a unique combination of tournament management, poker news, and engaging video content all in one platform.')
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'At a Glance')

    create_styled_table(doc, [
        ['Item', 'Details'],
        ['Delivery', 'Fully functional iOS & Android apps; Admin dashboard'],
        ['Timeline', '12 weeks from kickoff'],
        ['Ownership', 'You own all source code'],
        ['Support', '30-day post-launch warranty included'],
    ], col_widths=[2, 4.5])

    doc.add_paragraph()

    # ==================== 2. GOALS & OUTCOMES ====================
    add_section_header(doc, '2. Goals & Outcomes')

    add_subsection_header(doc, 'Our Vision')

    vision_items = [
        ('One-stop destination:', 'Your users can discover upcoming tournaments, watch highlight reels, and stay current with news—all within a single, polished mobile experience.'),
        ('Simple management:', 'A dedicated admin panel lets your team manage tournaments, articles, and videos without technical help.'),
        ('Scalable platform:', 'The architecture is built to grow with your business, supporting future feature additions seamlessly.'),
    ]

    for label, desc in vision_items:
        p = doc.add_paragraph()
        run = p.add_run(f'• {label} ')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'Who Benefits')

    create_styled_table(doc, [
        ['Audience', 'Key Benefits'],
        ['Players', 'Browse tournaments, view schedules & buy-ins, register interest'],
        ['Fans', 'Watch highlight videos, read poker news'],
        ['Organizers', 'Manage tournament info, publish news, maintain video library'],
    ], col_widths=[2, 4.5])

    doc.add_paragraph()

    # ==================== 3. WHAT YOU GET ====================
    add_section_header(doc, '3. What You Get')

    add_subsection_header(doc, 'A) Mobile Application (iOS & Android)')

    p = doc.add_paragraph()
    run = p.add_run('A native-feel app that works on both platforms:')
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    add_h3_header(doc, 'Home Screen')
    for item in ['Highlight video carousel with auto-play previews', 'Quick-access grid for Events, News, Videos, and more', 'Latest news headlines']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    add_h3_header(doc, 'Tournaments / Series')
    for item in ['List of upcoming, live, and completed events', 'Search, filter by date range, buy-in, game type', 'Detail pages with description, schedule, prize pool']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    add_h3_header(doc, 'News')
    for item in ['Article feed sorted by recency or category', 'Full article reader with images, author info']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    add_h3_header(doc, 'Videos')
    for item in ['Video gallery with thumbnails and durations', 'In-app playback (or YouTube redirect)']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    add_h3_header(doc, 'User Account')
    for item in ['Email/password authentication (email verification)', 'Profile screen with avatar and display name']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'B) Admin Dashboard (Web)')

    p = doc.add_paragraph()
    run = p.add_run('A secure management console for your team:')
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    admin_features = [
        ('Tournament Management', ['Create / edit / delete tournaments', 'Set dates, buy-ins, guarantee, blind structures', 'Toggle visibility, mark featured events']),
        ('News Management', ['Rich-text editor with image uploads', 'Assign categories, authors', 'Schedule or instant publish']),
        ('Video Management', ['Upload directly or paste YouTube/external links', 'Auto-extract thumbnail and duration', 'Sort videos into playlists']),
        ('User Management', ['View registered users, reset passwords', 'Assign admin or moderator roles']),
    ]

    for section_title, items in admin_features:
        add_h3_header(doc, section_title)
        for item in items:
            p = doc.add_paragraph()
            run = p.add_run(f'• {item}')
            run.font.size = Pt(11)
            run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'C) Backend & Infrastructure')

    create_styled_table(doc, [
        ['Component', 'Details'],
        ['API Server', 'Node.js / Express, secured with JWT'],
        ['Database', 'PostgreSQL (or MySQL—your choice)'],
        ['Media', 'Cloudinary for images/videos'],
        ['Email', 'SendGrid for transactional emails'],
    ], col_widths=[2, 4.5])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Everything is deployed on your cloud account for full ownership.')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = TEXT_COLOR

    doc.add_page_break()

    # ==================== 4. PROJECT TIMELINE ====================
    add_section_header(doc, '4. Project Timeline (12 Weeks)')

    create_styled_table(doc, [
        ['Stage', 'Duration', 'Key Outcomes'],
        ['Discovery & Setup', 'Weeks 1-2', 'Brand guidelines locked, cloud accounts provisioned, CI/CD ready'],
        ['Backend Build', 'Weeks 3-5', 'API endpoints for tournaments, news, videos, auth'],
        ['Admin Dashboard', 'Weeks 6-7', 'All CRUD screens, media uploads, role permissions'],
        ['Mobile App', 'Weeks 8-11', 'iOS & Android builds, navigation, API integration'],
        ['Testing & Launch', 'Week 12', 'Beta testing, bug fixes, App Store / Play Store submission'],
    ], col_widths=[2, 1.5, 3])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Our target is to submit to the app stores by Week 12, subject to platform review times.')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'Communication & Reporting')

    for item in ['Simple weekly progress report', 'Short demo every two weeks', 'Transparent issue tracking with real-time access']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    # ==================== 5. CAPACITY FIT ====================
    add_section_header(doc, '5. Capacity Fit')

    p = doc.add_paragraph()
    run = p.add_run("We're confident we can deliver a high-quality poker platform because:")
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    create_styled_table(doc, [
        ['Experience', 'Details'],
        ['API Design', 'We have built REST APIs serving thousands of concurrent users'],
        ['Mobile', 'Our team has launched 10+ apps on iOS and Android'],
        ['Admin Panels', 'We regularly build dashboards with Next.js / React'],
        ['Security', 'JWT authentication, role-based access, encrypted passwords'],
    ], col_widths=[2, 4.5])

    doc.add_paragraph()

    # ==================== 6. YOUR RESPONSIBILITIES ====================
    add_section_header(doc, '6. Your Responsibilities')

    add_subsection_header(doc, 'Simple Checklist')

    p = doc.add_paragraph()
    run = p.add_run('To ensure smooth delivery, we need you to:')
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    checklist = [
        'Provide brand assets (logo, fonts, color palette)',
        'Approve designs and key screens within 3 business days where possible',
        'Supply sample content (tournaments, articles, videos) for testing',
        'Test builds during the beta phase and share feedback',
        'Nominate a point-of-contact for weekly reviews',
    ]

    for item in checklist:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'Timely Feedback Required')

    p = doc.add_paragraph()
    run = p.add_run('This timeline requires your feedback within 3 business days on all deliverables. Delays in approval will extend the project timeline accordingly.')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    # ==================== 7. MAINTENANCE & SUPPORT ====================
    add_section_header(doc, '7. Maintenance & Support')

    p = doc.add_paragraph()
    run = p.add_run('Once launched, we offer two support tiers:')
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'Option A: Care+ 24/7')

    create_styled_table(doc, [
        ['Feature', 'Details'],
        ['Monthly Fee', 'RM 20,000/month'],
        ['Availability', '7 days a week, including public holidays'],
        ['Response', 'Critical issues within 2 hours, others within 8 hours'],
        ['Event Days', 'Up to 4 high-priority event days per month'],
        ['Improvements', '10 hours of minor enhancements / month'],
    ], col_widths=[2, 4.5])

    doc.add_paragraph()

    add_subsection_header(doc, 'Option B: Care Standard')

    create_styled_table(doc, [
        ['Feature', 'Details'],
        ['Monthly Fee', 'RM 11,000/month'],
        ['Availability', 'Monday – Friday, 9am – 6pm'],
        ['Response', 'Critical issues within 4 hours, others within 24 hours'],
        ['Event Days', 'Up to 1 high-priority event day per month'],
        ['Improvements', '4 hours of minor enhancements / month'],
    ], col_widths=[2, 4.5])

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('Both plans include hosting monitoring, security patches, and routine backups.')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = TEXT_COLOR

    doc.add_page_break()

    # ==================== 8. PRICING & PAYMENT ====================
    add_section_header(doc, '8. Pricing & Payment')

    add_subsection_header(doc, 'Investment Summary')

    create_styled_table(doc, [
        ['Item', 'Amount'],
        ['Total Development Fee', 'RM 250,000 (one-time)'],
        ['Monthly Support', 'RM 11,000 – RM 20,000 depending on option'],
    ], col_widths=[3, 3.5])

    doc.add_paragraph()

    add_subsection_header(doc, 'Payment Milestones')

    create_styled_table(doc, [
        ['Milestone', 'Timing', 'Amount', 'Percentage'],
        ['Contract Signing', 'Upon signing', 'RM 62,500', '25%'],
        ['Backend Complete', 'Week 5', 'RM 75,000', '30%'],
        ['Admin + Mobile Alpha', 'Week 9', 'RM 75,000', '30%'],
        ['Final Delivery', 'Week 12', 'RM 37,500', '15%'],
    ], col_widths=[2.2, 1.5, 1.5, 1.3])

    doc.add_paragraph()

    add_subsection_header(doc, "What's Included")

    included = [
        'Design & build of mobile app (iOS + Android)',
        'Backend API development',
        'Admin dashboard',
        'Cloudinary & SendGrid integration',
        'App Store and Play Store submission assistance',
        '4 hours of stakeholder training',
        '30-day post-launch warranty',
    ]

    for item in included:
        p = doc.add_paragraph()
        run = p.add_run(f'✓ {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = SUCCESS_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, "What's NOT Included")

    not_included = [
        'Apple Developer & Google Play annual fees',
        'Third-party service subscriptions after free-tier limits',
        'Content creation (articles, videos)',
        'Marketing or ASO services',
    ]

    for item in not_included:
        p = doc.add_paragraph()
        run = p.add_run(f'✗ {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = DANGER_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'Growth & Usage-Based Costs')

    p = doc.add_paragraph()
    run = p.add_run('As your user base grows, third-party provider charges (cloud compute, media storage, email/push sends) may increase. We will:')
    run.font.size = Pt(11)
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    for item in ['Enable spend alerts and budgets in your accounts', 'Provide a simple monthly forecast', 'Optimize caching/CDN to minimize costs']:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    # ==================== 9. DATA & PRIVACY ====================
    add_section_header(doc, '9. Data & Privacy')

    add_subsection_header(doc, 'Our Commitment')

    privacy_items = [
        'All user data is encrypted in transit and at rest',
        'Passwords are hashed using industry-standard algorithms',
        'Data is stored on servers you own; we do not retain copies',
        'You can export or delete all data at any time',
        'Local legal compliance (PDPA, GDPR) remains with you',
    ]

    for item in privacy_items:
        p = doc.add_paragraph()
        run = p.add_run(f'• {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    # ==================== 10. ACCEPTANCE & NEXT STEPS ====================
    add_section_header(doc, '10. Acceptance & Next Steps')

    add_subsection_header(doc, 'How to Get Started')

    steps = [
        ('Step 1:', 'Sign this proposal and pay the first milestone (25%)'),
        ('Step 2:', 'Kickoff call (1 hour) to confirm brand assets, timelines, and communication channels'),
        ('Step 3:', 'We begin work within 5 business days of kickoff'),
    ]

    for step, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(step + ' ')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = ACCENT_COLOR
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_page_break()

    # ==================== SIGNATURES ====================
    add_section_header(doc, 'Signatures')

    add_subsection_header(doc, 'Client')

    sig_fields = ['Name:', 'Title:', 'Signature:', 'Date:']

    for field in sig_fields:
        p = doc.add_paragraph()
        run = p.add_run(field)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        run = p.add_run(' ' + '_' * 50)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(14)

    doc.add_paragraph()

    add_subsection_header(doc, 'Service Provider')

    for field in sig_fields:
        p = doc.add_paragraph()
        run = p.add_run(field)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        run = p.add_run(' ' + '_' * 50)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('This proposal is confidential and valid for 30 days.')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Thank you for considering Poker Dream.')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    # Save document
    output_path = '/Users/clifflai/development/vsc-workspace/poker-dream/POKER_DREAM_PROPOSAL.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    main()
