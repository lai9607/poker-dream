#!/usr/bin/env python3
"""
Generate a professionally designed Word document for Poker Dream Website Development Proposal
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Brand Colors
PRIMARY_COLOR = RGBColor(0x2C, 0x3E, 0x50)      # Dark Blue
SECONDARY_COLOR = RGBColor(0x34, 0x49, 0x5E)   # Slate
ACCENT_COLOR = RGBColor(0xC0, 0xA0, 0x62)      # Gold
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)        # Dark Gray
SUCCESS_COLOR = RGBColor(0x27, 0xAE, 0x60)     # Green
DANGER_COLOR = RGBColor(0xC0, 0x39, 0x2B)      # Red

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Add a horizontal line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C0A062')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_styled_table(doc, data, header=True, col_widths=None):
    """Create a professionally styled table"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    if i == 0 and header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            if i == 0 and header:
                set_cell_shading(cell, '2C3E50')
            elif i % 2 == 0:
                set_cell_shading(cell, 'F8F9FA')

    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    return table

def add_section_header(doc, title):
    """Add a section header"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    add_horizontal_line(doc)

def add_subsection_header(doc, title):
    """Add a subsection header"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR

def add_h3_header(doc, title):
    """Add a h3 header"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_COLOR

def add_bullet_point(doc, text, color=TEXT_COLOR):
    """Add a bullet point"""
    p = doc.add_paragraph()
    run = p.add_run(f'• {text}')
    run.font.size = Pt(11)
    run.font.color.rgb = color

def add_paragraph_text(doc, text, bold=False, italic=False, color=TEXT_COLOR):
    """Add paragraph text"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color

def main():
    doc = Document()

    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ==================== COVER PAGE ====================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BUSINESS PROPOSAL')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Poker Dream Website')
    run.font.size = Pt(28)
    run.font.color.rgb = PRIMARY_COLOR

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Website Development')
    run.font.size = Pt(20)
    run.font.color.rgb = SECONDARY_COLOR

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"Your Poker Community Online"')
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = ACCENT_COLOR

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_data = [
        ['Version:', '1.0'],
        ['Date:', '12 December 2025'],
        ['Validity:', '30 days (until 12 January 2026)'],
    ]

    for label, value in info_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label} ')
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = TEXT_COLOR
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.color.rgb = TEXT_COLOR

    doc.add_page_break()

    # ==================== EXECUTIVE SUMMARY ====================
    add_section_header(doc, 'EXECUTIVE SUMMARY')

    add_paragraph_text(doc,
        'We are pleased to present this proposal for the development of the Poker Dream Website - '
        'a responsive, modern website designed to serve as the public-facing platform for tournament '
        'information, player rankings, news, and brand presence, integrated with your existing backend infrastructure.')

    doc.add_paragraph()

    add_subsection_header(doc, 'The Project')

    create_styled_table(doc, [
        ['Project', 'Description'],
        ['Poker Dream Website', 'Public-facing responsive website modeled after pokerdream-live.com, '
         'featuring tournament information, DPOY rankings, news, galleries, and integrated with existing backend APIs'],
    ], col_widths=[2.5, 4])

    doc.add_paragraph()

    add_subsection_header(doc, 'Development Approach')

    add_paragraph_text(doc,
        'This proposal outlines a custom Next.js website development approach that integrates directly '
        'with your existing Poker Dream backend APIs. The website will share data models with your Flutter '
        'mobile app, ensuring consistency across all platforms.')

    doc.add_page_break()

    # ==================== PROJECT SCOPE ====================
    add_section_header(doc, 'PROJECT SCOPE')

    add_subsection_header(doc, 'Website Analysis: pokerdream-live.com')

    add_paragraph_text(doc,
        'Based on our analysis of the reference website (pokerdream-live.com), we have identified '
        '14-16 main pages with the following features:')

    doc.add_paragraph()

    create_styled_table(doc, [
        ['Page', 'Features'],
        ['Home', 'Hero section, Featured tournaments, Winners gallery carousel, Instagram feed, Newsletter signup'],
        ['Tournament (List)', 'Searchable/filterable tournament listing with pagination, Status indicators (Live/Completed/Upcoming)'],
        ['Tournament (Detail)', 'Live timer, Blind levels, Chip counts, Payout structures, Player standings'],
        ['DPOY (Leaderboard)', 'Player of the Year rankings, Points system, Pagination (50+ entries per page)'],
        ['Live Reports', 'Tournament updates hub with event cards'],
        ['Gallery - Tournament', 'Photo gallery organized by tournament series (PD10-PD18+)'],
        ['Gallery - Champion', 'Winner photos grid with 200+ images'],
        ['News/Events', 'Promotions and announcements grid with article cards'],
        ['About Us', 'Company information, Mission, Team credentials'],
        ['Contact Us', 'Partnership inquiry form, Social links'],
        ['Career', 'Job listings (7 positions) with detailed descriptions'],
        ['Privacy Policy', 'Legal documentation'],
        ['Terms of Service', 'Legal documentation'],
    ], col_widths=[2, 4.5])

    doc.add_paragraph()

    add_subsection_header(doc, 'Key Features to Implement')

    create_styled_table(doc, [
        ['Module', 'Features'],
        ['Real-time Tournament Tracking', 'Live blind level timer with countdown, Current average chip count, Prize pool calculations, Payout structure display'],
        ['Player Ranking System (DPOY)', 'Points-based leaderboard, Multi-page pagination, Country flags and player profiles'],
        ['Content Management', 'News/Events publishing, Tournament schedules, Photo galleries (Tournament & Champions)'],
        ['User Engagement', 'Newsletter subscription, Social media integration (FB, IG, Twitter, WhatsApp, YouTube), Partnership inquiry forms'],
        ['Multi-Language Support', 'English and Chinese player guides and content'],
        ['SEO & Performance', 'Server-side rendering, Meta tags optimization, Google Analytics integration'],
    ], col_widths=[2.5, 4])

    doc.add_page_break()

    # ==================== TECHNICAL APPROACH ====================
    add_section_header(doc, 'TECHNICAL APPROACH')

    add_subsection_header(doc, 'Recommended: Custom Next.js Development')

    add_paragraph_text(doc,
        'Given you already have a Flutter app with structured data models and existing backend APIs, '
        'we recommend building a custom Next.js website for the following reasons:')

    doc.add_paragraph()

    benefits = [
        ('Direct API Integration', 'Seamlessly connects with your existing backend infrastructure'),
        ('Shared Data Models', 'Tournament, player, and content data synchronized across mobile app and website'),
        ('Superior SEO', 'Server-side rendering for better search engine visibility'),
        ('Real-time Features', 'Native support for WebSocket connections and live updates'),
        ('Performance', 'Optimized loading with automatic code splitting and image optimization'),
        ('Maintainability', 'Modern React codebase that your team can extend and maintain'),
    ]

    for label, desc in benefits:
        p = doc.add_paragraph()
        run = p.add_run(f'• {label}: ')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, 'Technology Stack')

    create_styled_table(doc, [
        ['Component', 'Technology'],
        ['Framework', 'Next.js 14 (React)'],
        ['Language', 'TypeScript'],
        ['Styling', 'Tailwind CSS'],
        ['State Management', 'React Query / SWR'],
        ['Hosting', 'Vercel or AWS'],
        ['CDN', 'Cloudflare'],
        ['Analytics', 'Google Analytics 4'],
    ], col_widths=[2.5, 4])

    doc.add_page_break()

    # ==================== TIMELINE ====================
    add_section_header(doc, 'PROJECT TIMELINE')

    add_subsection_header(doc, 'Website Development')

    p = doc.add_paragraph()
    run = p.add_run('Duration: 13-17 weeks')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    doc.add_paragraph()

    add_paragraph_text(doc,
        'Objective: Launch a fully responsive website integrated with existing backend APIs, '
        'featuring tournament information, player rankings, news, and content management.')

    doc.add_paragraph()

    add_subsection_header(doc, 'Timeline Breakdown')

    create_styled_table(doc, [
        ['Component', 'Duration', 'Description'],
        ['Discovery & Design', '2 weeks', 'Requirements gathering, UI/UX design, wireframes, design system'],
        ['Foundation & Core Pages', '3-4 weeks', 'Project setup, component library, Home, About, Contact, Career, Legal pages'],
        ['Tournament System', '4-5 weeks', 'Tournament list/detail pages, real-time timer, payout structures, standings'],
        ['Player & Rankings', '2-3 weeks', 'DPOY leaderboard, player profiles, pagination'],
        ['Content & Media', '2-3 weeks', 'News/Events, Live Reports, Tournament & Champion galleries'],
        ['Integration & Launch', '2 weeks', 'Newsletter, social media, SEO, testing, deployment'],
    ], col_widths=[2, 1.5, 3])

    doc.add_paragraph()

    add_subsection_header(doc, 'Deliverables - Website')

    create_styled_table(doc, [
        ['Module', 'Features Included'],
        ['Home Page', 'Hero section with featured tournaments, Winners gallery carousel, Instagram feed integration, Newsletter signup, Command grid navigation'],
        ['Tournament Pages', 'Searchable/filterable listing, Detail pages with live timer, Blind level countdown, Chip counts, Payout structures, Player standings'],
        ['DPOY Leaderboard', 'Points-based rankings, Pagination, Country flags, Player statistics'],
        ['News & Events', 'Article listing with filters, Full article detail pages, Promotions display'],
        ['Live Reports', 'Tournament update cards, Event-specific pages'],
        ['Galleries', 'Tournament photo gallery by series, Champion gallery with 200+ images, Lightbox viewing'],
        ['Information Pages', 'About Us, Contact Us with partnership form, Career page with job listings'],
        ['Legal Pages', 'Privacy Policy, Terms of Service'],
    ], col_widths=[2, 4.5])

    doc.add_paragraph()

    add_subsection_header(doc, 'Additional Deliverables')

    create_styled_table(doc, [
        ['Deliverable', 'Description'],
        ['Design System', 'Figma files, component library, brand guidelines'],
        ['Technical Documentation', 'Architecture docs, API integration specs, deployment guides'],
        ['SEO Setup', 'Meta tags, sitemap, robots.txt, structured data'],
        ['Analytics', 'Google Analytics 4 integration with custom events'],
        ['DevOps', 'CI/CD pipeline, staging/production environments'],
    ], col_widths=[2.5, 4])

    doc.add_page_break()

    # ==================== CAPACITY FIT ====================
    add_section_header(doc, 'CAPACITY FIT')

    add_paragraph_text(doc, 'Your expected scale:')

    doc.add_paragraph()

    create_styled_table(doc, [
        ['Metric', 'Your Estimate'],
        ['Admin Users', '~5 staff managing content'],
        ['Content Updates', 'Tournament updates during events, news weekly'],
        ['Monthly Visitors', '~50,000 - 100,000 page views'],
    ], col_widths=[3, 3.5])

    doc.add_paragraph()

    add_paragraph_text(doc,
        'This architecture comfortably supports your scale, with headroom for traffic spikes during '
        'major events. The CDN and caching strategy ensures fast load times globally.')

    doc.add_page_break()

    # ==================== INVESTMENT SUMMARY ====================
    add_section_header(doc, 'INVESTMENT SUMMARY')

    add_subsection_header(doc, 'Total Project Investment')

    create_styled_table(doc, [
        ['Phase', 'Duration', 'Investment'],
        ['Website Development', '13-17 weeks', 'RM 190,640'],
    ], col_widths=[2.5, 2, 2])

    doc.add_paragraph()

    add_subsection_header(doc, 'Cost Breakdown by Category')

    create_styled_table(doc, [
        ['Category', 'Amount'],
        ['Discovery & Design', 'RM 22,000'],
        ['Foundation & Core Pages', 'RM 39,440'],
        ['Tournament System', 'RM 42,160'],
        ['Player & Rankings', 'RM 23,120'],
        ['Content & Media', 'RM 29,920'],
        ['Integration & Launch', 'RM 34,000'],
        ['Total', 'RM 190,640'],
    ], col_widths=[3, 3.5])

    doc.add_paragraph()

    add_subsection_header(doc, 'Payment Schedule')

    create_styled_table(doc, [
        ['Milestone', 'Percentage', 'Amount'],
        ['Contract Signing', '50%', 'RM 95,320'],
        ['Upon Deliver Core Pages & Tournament List', '20%', 'RM 38,128'],
        ['Upon Deliver DPOY & Content Pages', '15%', 'RM 28,596'],
        ['Project Completion', '15%', 'RM 28,596'],
        ['TOTAL', '100%', 'RM 190,640'],
    ], col_widths=[3, 1.5, 2])

    doc.add_paragraph()

    add_subsection_header(doc, "What's Included")

    included = [
        'Design & build of responsive website (Desktop, Tablet, Mobile)',
        'Next.js frontend development with TypeScript',
        'Integration with existing backend APIs',
        'SEO optimization and Google Analytics setup',
        'CDN and hosting configuration',
        'Newsletter integration',
        'Social media feed integration',
        '2 hours of stakeholder training',
        '30-day post-launch warranty',
    ]

    for item in included:
        p = doc.add_paragraph()
        run = p.add_run(f'✓ {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = SUCCESS_COLOR

    doc.add_paragraph()

    add_subsection_header(doc, "What's NOT Included")

    not_included = [
        'Backend API development (using existing APIs)',
        'Content creation (articles, photos, videos)',
        'Domain registration and annual renewal',
        'Third-party service subscriptions (email service, etc.)',
        'Marketing or SEO services beyond initial setup',
    ]

    for item in not_included:
        p = doc.add_paragraph()
        run = p.add_run(f'✗ {item}')
        run.font.size = Pt(11)
        run.font.color.rgb = DANGER_COLOR

    doc.add_page_break()

    # ==================== MONTHLY RECURRING COSTS ====================
    add_section_header(doc, 'MONTHLY RECURRING COSTS')

    add_paragraph_text(doc, 'Estimated monthly costs for hosting and third-party services:')

    doc.add_paragraph()

    create_styled_table(doc, [
        ['Service', 'Est. Monthly Cost'],
        ['Hosting (Vercel/AWS)', 'RM 225 - RM 900'],
        ['CDN (Cloudflare)', 'RM 0 - RM 225'],
        ['Email Service (Newsletter)', 'RM 90 - RM 225'],
        ['Domain Renewal', '~RM 70/year'],
        ['SSL Certificate', 'Free (Let\'s Encrypt)'],
        ['Estimated Monthly Total', 'RM 315 - RM 1,350'],
    ], col_widths=[3, 3.5])

    doc.add_paragraph()

    add_paragraph_text(doc,
        'Costs will vary based on traffic and usage. We will set up monitoring and alerts '
        'to help you track and optimize these costs.',
        italic=True)

    doc.add_page_break()

    # ==================== POST-LAUNCH SUPPORT ====================
    add_section_header(doc, 'POST-LAUNCH SUPPORT PACKAGES')

    add_paragraph_text(doc,
        'After launch, we offer ongoing support packages to ensure your website remains '
        'stable, secure, and up-to-date:')

    doc.add_paragraph()

    create_styled_table(doc, [
        ['Package', 'Name', 'Monthly Investment', 'Includes'],
        ['Option A', 'Essential', 'RM 11,000/month', 'Bug fixes, security patches, minor updates, 8 hours support'],
        ['Option B', 'Growth', 'RM 20,000/month', 'Essential + feature enhancements, performance optimization, 20 hours support'],
        ['Option C', 'Enterprise', 'RM 35,000/month', 'Growth + dedicated developer, priority support, 40 hours support, SLA'],
    ], col_widths=[1, 1.2, 1.8, 2.5])

    doc.add_paragraph()

    add_subsection_header(doc, 'Option A: Essential Support')

    p = doc.add_paragraph()
    run = p.add_run('Monthly Fee: RM 11,000/month')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    doc.add_paragraph()

    add_paragraph_text(doc, 'Designed for: Regular operations with standard content updates.')

    essential_items = [
        'Bug fixes and security patches',
        'Monthly security & reliability upkeep',
        'Up to 8 hours/month for minor improvements',
        'Business hours support (Mon-Fri, 9am-6pm)',
        'Response time: Critical within 4 hours, others within 24 hours',
    ]

    for item in essential_items:
        add_bullet_point(doc, item)

    doc.add_paragraph()

    add_subsection_header(doc, 'Option B: Growth Support')

    p = doc.add_paragraph()
    run = p.add_run('Monthly Fee: RM 20,000/month')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    doc.add_paragraph()

    add_paragraph_text(doc, 'Designed for: Active feature development and frequent updates.')

    growth_items = [
        'Everything in Essential',
        'Feature enhancements and optimizations',
        'Up to 20 hours/month for improvements',
        '24/7 on-call coverage',
        'Response time: Critical within 2 hours',
        'Up to 4 event days/month remote support',
    ]

    for item in growth_items:
        add_bullet_point(doc, item)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run('*Recommended: Growth package for active development and tournament event support.')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_page_break()

    # ==================== TERMS & CONDITIONS ====================
    add_section_header(doc, 'TERMS & CONDITIONS')

    terms = [
        ('Payment Terms', 'Net 15 days from invoice date. Late payments may incur a 2% monthly interest charge.'),
        ('Milestone Acceptance', '5 business days review period for each milestone. Acceptance is implied if no feedback is received within the review period.'),
        ('Change Requests', 'Any changes to the agreed scope will be handled via a formal change order process. Additional costs and timeline adjustments will be quoted separately.'),
        ('Intellectual Property', 'Full IP ownership transfers to the client upon final payment. Source code, design files, and all assets become client property.'),
        ('Confidentiality', 'A Non-Disclosure Agreement (NDA) is included with the contract. Both parties agree to keep project details confidential.'),
        ('Warranty', '30-day bug fix warranty post-launch. Covers defects in the delivered functionality as per agreed specifications.'),
        ('Third-Party Services', 'Costs for third-party services (hosting, CDN, analytics tools) are not included and are the client\'s responsibility.'),
        ('Project Cancellation', 'If the project is cancelled, payment is due for all completed work and work in progress up to the cancellation date.'),
        ('Communication', 'Weekly progress updates via video call. Day-to-day communication via Slack or agreed platform.'),
        ('Timeline', 'Timelines are estimates based on current scope. Delays caused by client feedback or scope changes may extend the timeline.'),
    ]

    for title, desc in terms:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_COLOR

        p = doc.add_paragraph()
        run = p.add_run(desc)
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT_COLOR
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ==================== CLOSING ====================
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('This proposal is confidential and intended solely for the named recipient.')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Pricing is valid for 30 days from the proposal date.')
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = TEXT_COLOR

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Prepared with confidence. Ready to build the home for poker.')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = ACCENT_COLOR

    # Save document
    output_path = '/Users/clifflai/development/vsc-workspace/poker-dream/Poker_Dream_Website_Proposal.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    main()
